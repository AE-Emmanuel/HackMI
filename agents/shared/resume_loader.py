"""Multi-format resume text extraction.

Single entry point: :func:`load_resume(path)` returns the plain-text body
of a resume file regardless of source format. Supported formats:

- ``.txt`` / ``.md``  — direct read.
- ``.docx``           — via ``python-docx`` (paragraphs + tables).
- ``.pdf``            — via ``pdfplumber`` (text-extractable PDFs only).
- ``.rtf``            — best-effort regex strip of RTF control words.
- ``.doc``            — legacy binary Word; not supported (raise clearly).

Bytes-in API (:func:`extract_resume_bytes`) lets the FastAPI ``/analyze``
endpoint accept uploads directly without writing them to disk.

Whitespace is collapsed to single spaces within lines, blank lines are
preserved as paragraph separators. Empty extractions raise
:class:`ResumeExtractionError` so the pipeline fails fast rather than
feeding garbage into Agent 1.
"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import IO


class ResumeExtractionError(ValueError):
    """Raised when a resume file cannot be read or yields no usable text."""


# ---------------------------------------------------------------------------
# Format-specific extractors (operate on file-like bytes objects)
# ---------------------------------------------------------------------------


def _extract_txt(stream: IO[bytes]) -> str:
    raw = stream.read()
    # Try UTF-8 first, fall back to latin-1 (lossless for any byte sequence).
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _extract_docx(stream: IO[bytes]) -> str:
    try:
        from docx import Document  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise ResumeExtractionError(
            "python-docx is required for .docx files. Install it with "
            "`uv pip install python-docx`."
        ) from exc

    doc = Document(stream)
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Pull text from any tables (some resumes use tables for layout).
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def _extract_pdf(stream: IO[bytes]) -> str:
    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover
        raise ResumeExtractionError(
            "pdfplumber is required for .pdf files. Install it with "
            "`uv pip install pdfplumber`."
        ) from exc

    parts: list[str] = []
    with pdfplumber.open(stream) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)
    return "\n".join(parts)


_RTF_CONTROL_RE = re.compile(r"\\[a-zA-Z]+-?\d*\s?|\\\W|\{|\}")


def _extract_rtf(stream: IO[bytes]) -> str:
    raw = stream.read().decode("latin-1", errors="replace")
    # Strip RTF control words and braces — crude but works for plain resumes.
    cleaned = _RTF_CONTROL_RE.sub(" ", raw)
    return re.sub(r"\s+", " ", cleaned).strip()


_EXTRACTORS = {
    ".txt": _extract_txt,
    ".md": _extract_txt,
    ".markdown": _extract_txt,
    ".docx": _extract_docx,
    ".pdf": _extract_pdf,
    ".rtf": _extract_rtf,
}


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def supported_extensions() -> list[str]:
    """List of supported file extensions (lowercase, with leading dot)."""

    return sorted(_EXTRACTORS.keys())


def _normalize_whitespace(text: str) -> str:
    """Collapse intra-line whitespace, preserve paragraph breaks."""

    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    # Collapse runs of blank lines into a single blank.
    out: list[str] = []
    last_blank = False
    for line in lines:
        if not line:
            if not last_blank:
                out.append("")
            last_blank = True
        else:
            out.append(line)
            last_blank = False
    return "\n".join(out).strip()


def extract_resume_bytes(data: bytes, filename: str) -> str:
    """Extract resume text from raw bytes given the original filename.

    Use this for HTTP uploads where you don't want to write to disk.
    """

    suffix = Path(filename).suffix.lower()
    if suffix == ".doc":
        raise ResumeExtractionError(
            "Legacy .doc (Word 97-2003 binary) is not supported. "
            "Open the file in Word and save as .docx."
        )
    extractor = _EXTRACTORS.get(suffix)
    if extractor is None:
        raise ResumeExtractionError(
            f"Unsupported resume format: {suffix!r}. "
            f"Supported: {', '.join(supported_extensions())}."
        )

    try:
        text = extractor(io.BytesIO(data))
    except ResumeExtractionError:
        raise
    except Exception as exc:
        raise ResumeExtractionError(
            f"Failed to extract text from {filename}: {exc}"
        ) from exc

    text = _normalize_whitespace(text)
    if not text:
        raise ResumeExtractionError(
            f"Extracted no text from {filename}. The file may be empty, "
            "image-only (scanned), or corrupted."
        )
    return text


def load_resume(path: str | Path) -> str:
    """Read a resume from disk and return its plain-text body."""

    p = Path(path)
    if not p.exists():
        raise ResumeExtractionError(f"Resume file not found: {p}")
    if not p.is_file():
        raise ResumeExtractionError(f"Not a file: {p}")

    return extract_resume_bytes(p.read_bytes(), p.name)
